import os
from io import BytesIO
from flask import send_file, jsonify, make_response
from docx import Document
from fpdf import FPDF
import arabic_reshaper
from bidi.algorithm import get_display
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

def format_text(text, language):
    """Format text for proper display based on language."""
    if language == "Arabic":
        return reshape_arabic(text)
    return text

def export_to_excel(chat_data):
    """Export chat data to Excel format. Requires pandas and openpyxl."""
    try:
        # Optional import - only works if packages are installed
        import pandas as pd
        
        # Prepare data for Excel
        data = []
        for item in chat_data:
            data.append({
                'Original Text': item.get('original', ''),
                'Translated Text': item.get('translated', ''),
                'Source Language': item.get('source_lang', ''),
                'Target Language': item.get('target_lang', '')
            })
        
        # Create DataFrame
        df = pd.DataFrame(data)
        
        # Create Excel file in memory
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Translations', index=False)
        
        output.seek(0)
        
        response = make_response(send_file(
            output,
            as_attachment=True,
            download_name='chat_translations.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        ))
        
        return response
        
    except ImportError:
        return jsonify({
            'error': 'Excel export not available. Install pandas and openpyxl for Excel export functionality.',
            'message': 'Run: pip install pandas openpyxl'
        }), 400
    except Exception as e:
        return jsonify({'error': f'Failed to export to Excel: {str(e)}'}), 500

def export_to_word(chat_data):
    """Export chat data to Word document format."""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('Translation Results', 0)
        title.alignment = 1  # Center alignment
        
        # Add content
        for i, item in enumerate(chat_data, 1):
            # Add entry number
            doc.add_heading(f'Translation {i}', level=1)
            
            # Original text
            original_para = doc.add_paragraph()
            original_para.add_run('Original: ').bold = True
            original_para.add_run(item.get('original', ''))
            
            # Translated text
            translated_para = doc.add_paragraph()
            translated_para.add_run('Translated: ').bold = True
            translated_para.add_run(item.get('translated', ''))
            
            # Languages
            lang_para = doc.add_paragraph()
            lang_para.add_run('Language Pair: ').bold = True
            lang_para.add_run(f"{item.get('source_lang', '')} → {item.get('target_lang', '')}")
            
            doc.add_paragraph('─' * 50)  # Separator
        
        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        response = make_response(send_file(
            output,
            as_attachment=True,
            download_name='chat_translations.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ))
        
        return response
        
    except Exception as e:
        return jsonify({'error': f'Failed to export to Word: {str(e)}'}), 500

def export_to_pdf_table(chat_data):
    """Export chat data to PDF with table format."""
    try:
        print(f"[DEBUG] Received chat_data: {chat_data}")  # Debug log
        print(f"[DEBUG] Chat data type: {type(chat_data)}")  # Debug log
        print(f"[DEBUG] Chat data length: {len(chat_data) if chat_data else 0}")  # Debug log
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        
        # Prepare data for table
        data = [['Original Text', 'Translated Text', 'Language Pair']]
        
        for i, item in enumerate(chat_data):
            print(f"[DEBUG] Processing item {i}: {item}")  # Debug log
            original = format_text(item.get('original', ''), item.get('source_lang', ''))
            translated = format_text(item.get('translated', ''), item.get('target_lang', ''))
            lang_pair = f"{item.get('source_lang', '')} → {item.get('target_lang', '')}"
            
            data.append([original, translated, lang_pair])
        
        # Create table
        table = Table(data, colWidths=[150, 150, 100])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        # Build PDF
        doc.build([table])
        buffer.seek(0)
        
        response = make_response(send_file(
            buffer,
            as_attachment=True,
            download_name='chat_translations.pdf',
            mimetype='application/pdf'
        ))
        
        return response
        
    except Exception as e:
        return jsonify({'error': f'Failed to export to PDF: {str(e)}'}), 500

def export_to_word_text(chat_data):
    """Export chat data to Word document with text format."""
    return export_to_word(chat_data)  # Same implementation

def export_to_pdf_text(chat_data):
    """Export chat data to PDF with text format."""
    try:
        print(f"[DEBUG] PDF Text - Received chat_data: {chat_data}")  # Debug log
        print(f"[DEBUG] PDF Text - Chat data type: {type(chat_data)}")  # Debug log
        print(f"[DEBUG] PDF Text - Chat data length: {len(chat_data) if chat_data else 0}")  # Debug log
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title = Paragraph("Translation Results", styles['Title'])
        story.append(title)
        
        # Add content
        for i, item in enumerate(chat_data, 1):
            print(f"[DEBUG] PDF Text - Processing item {i}: {item}")  # Debug log
            # Entry header
            header = Paragraph(f"Translation {i}", styles['Heading1'])
            story.append(header)
            
            # Original text
            original_text = f"<b>Original:</b> {item.get('original', '')}"
            original_para = Paragraph(original_text, styles['Normal'])
            story.append(original_para)
            
            # Translated text
            translated_text = f"<b>Translated:</b> {item.get('translated', '')}"
            translated_para = Paragraph(translated_text, styles['Normal'])
            story.append(translated_para)
            
            # Language pair
            lang_text = f"<b>Language Pair:</b> {item.get('source_lang', '')} → {item.get('target_lang', '')}"
            lang_para = Paragraph(lang_text, styles['Normal'])
            story.append(lang_para)
            
            # Separator
            separator = Paragraph("─" * 50, styles['Normal'])
            story.append(separator)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        response = make_response(send_file(
            buffer,
            as_attachment=True,
            download_name='chat_translations.pdf',
            mimetype='application/pdf'
        ))
        
        return response
        
    except Exception as e:
        return jsonify({'error': f'Failed to export to PDF: {str(e)}'}), 500

# Register arial font
font_path = os.path.join(os.path.dirname(__file__), 'fonts', 'arial.ttf')
if os.path.exists(font_path):
    pdfmetrics.registerFont(TTFont('arial', font_path))

def reshape_arabic(text):
    """Fix Arabic text display issues (inversion & spacing)."""
    return get_display(arabic_reshaper.reshape(text))